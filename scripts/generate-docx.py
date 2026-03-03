#!/usr/bin/env python3
"""Generate Source AI — comprehensive document explaining what Source does."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style setup ──────────────────────────────────────────────────────────────

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.bold = True

doc.styles["Heading 1"].font.size = Pt(20)
doc.styles["Heading 1"].paragraph_format.space_before = Pt(28)
doc.styles["Heading 1"].paragraph_format.space_after = Pt(10)

doc.styles["Heading 2"].font.size = Pt(15)
doc.styles["Heading 2"].paragraph_format.space_before = Pt(20)
doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)

doc.styles["Heading 3"].font.size = Pt(12)
doc.styles["Heading 3"].paragraph_format.space_before = Pt(14)
doc.styles["Heading 3"].paragraph_format.space_after = Pt(4)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1A1A1A")
        shading.set(qn("w:val"), "clear")
        cell._element.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            if r_idx % 2 == 1:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F5F5F5")
                shading.set(qn("w:val"), "clear")
                cell._element.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    return p


def bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        p.add_run(text).font.size = Pt(10.5)
    else:
        p.add_run(text).font.size = Pt(10.5)
    return p


def quote(text, attribution=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"\u201c{text}\u201d")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if attribution:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        run2 = p2.add_run(f"\u2014 {attribution}")
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ═════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════

for _ in range(7):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SOURCE AI")
run.bold = True
run.font.size = Pt(36)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("The Migration Engine for Consultancies")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "An AI engine that automates 80% of ERP migrations \u2014 so your consultancy\n"
    "wins more deals, moves faster, and makes more money."
)
run.font.size = Pt(11)
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("source.ai")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# THE PROBLEM
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("The Problem: Why ERP Migrations Are Broken", level=1)

para(
    "Mid-market companies \u2014 typically in the $10M\u2013$300M revenue range \u2014 eventually outgrow "
    "their starter accounting systems. QuickBooks, Xero, MYOB, Sage \u2014 these tools work well "
    "for small businesses, but they break down when a company scales. The chart of accounts gets "
    "unwieldy, inventory sync becomes manual, month-end close takes 12+ days, and international "
    "expansion is impossible. At that point, they need to migrate to a modern ERP like Oracle "
    "NetSuite or Microsoft Dynamics 365."
)

para(
    "The problem is that the migration itself is painful, expensive, and slow. The consulting "
    "industry that has grown up around ERP migrations follows a delivery chain that hasn\u2019t "
    "fundamentally changed in decades. A typical mid-market QuickBooks-to-NetSuite migration "
    "passes through six stages, each adding time, cost, and handoffs:"
)

add_table(
    ["Stage", "What Happens"],
    [
        ["1. End Client", "Identifies the need, issues an RFP, evaluates vendors"],
        ["2. Sales & Scoping", "Discovery calls, needs analysis, proposal writing, SOW negotiation"],
        ["3. Solutions Architecture", "Schema analysis, gap assessment, AS-IS \u2192 TO-BE mapping, BRD creation"],
        ["4. Resource & Team Build", "Find offshore/onshore developers, onboard them, ramp up on client systems"],
        ["5. Implementation & Config", "Data migration, custom development, integrations, UAT, bug fixing"],
        ["6. Go-Live & Handover", "Cutover planning, validation, training, hypercare, documentation"],
    ],
    col_widths=[1.8, 4.9],
)

para(
    "The typical timeline is 4\u20136+ months. The typical cost to the end client is $30K\u2013$40K in "
    "consulting fees alone, on top of the ERP license itself. And the consulting firm doing the "
    "work doesn\u2019t make nearly as much as people assume. For every $200/hr billed, roughly $150 "
    "gets consumed by marketing and sales costs to win the deal, project management overhead, "
    "offshore team costs, customer success, and implementation chaos. Typical mid-market "
    "consultancy margins on migrations sit at just 20\u201330%. On a $30K migration, the firm might "
    "net roughly $6K in actual profit."
)

para(
    "Being a mid-market ERP consultant is, frankly, a grind. Winning net-new migration deals "
    "means competing in bake-offs with four or five other firms, managing post-migration chaos, "
    "handling change orders, and dealing with frustrated customers who feel like the project took "
    "too long, cost too much, and involved too many revisions. Migrations are thin-margin, "
    "time-intensive work \u2014 and both the consultant and the end client are frustrated by the process."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# WHAT SOURCE AI IS
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("What Source AI Is", level=1)

para(
    "Source AI is an autonomous AI engine that automates 80% of ERP migration work for "
    "mid-market companies. It is not a tool that consultants log into \u2014 it is a service that "
    "replaces the expensive, slow middle of the delivery chain. Source handles the solutions "
    "architecture, the offshoring, the implementation, and the migration execution. The "
    "consultant keeps the client relationship, the sales, the scoping, and the final review."
)

para(
    "The way to think about it is this: today, when a consultancy wins a migration deal, they "
    "need to spin up a solutions architect, possibly an offshore development team, and spend "
    "months going through BRD cycles, data mapping, configuration, and testing. With Source, "
    "the consultant\u2019s sales person goes on the discovery call, captures the context, and then "
    "hands it off to Source. Source connects to the client\u2019s existing system, scans it, infers "
    "the business logic, generates the BRD, creates the migration plan, executes the "
    "implementation, and delivers a live, configured target system \u2014 all in under 21 days."
)

para(
    "Source is not trying to replace consultancies. It partners with them. The consultant "
    "remains the face of the engagement. The client never sees Source\u2019s branding \u2014 every "
    "deliverable, every questionnaire, every report ships under the consultancy\u2019s own logo, "
    "domain, and colours. Source is the engine under the hood that makes the consultancy "
    "dramatically more efficient and profitable."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# HOW IT WORKS
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("How Source AI Works", level=1)

doc.add_heading("The Old Delivery Chain vs. The New Chain", level=2)

para(
    "In the traditional model, the delivery chain runs: End Client \u2192 Sales \u2192 Solutions "
    "Architect \u2192 Offshore Team \u2192 Migration Complete. With Source, the chain collapses to: "
    "End Client \u2192 Your Consultant \u2192 Source AI Engine \u2192 Live System. The Solutions Architect "
    "and the offshore team are eliminated entirely. The consultant stays on the relationship "
    "and the commercial ownership. Source does the rest."
)

doc.add_heading("What Source AI Handles Autonomously", level=2)

para(
    "Source takes on the full technical execution of a migration \u2014 approximately 85% of the "
    "total work. This includes:"
)

capabilities = [
    ("System Scanning", "Connects to the client\u2019s legacy system (QuickBooks, Xero, MYOB, etc.) via API with read-only access. Builds a complete instance map \u2014 schemas, fields, customisations, data volumes, integrations."),
    ("Business Logic Inference", "Analyses the scanned data to understand what\u2019s actually happening in the system versus what was described. Identifies workflows, customisations, and tribal knowledge embedded in the data."),
    ("Strategy Generation", "Identifies pain points, business workflows, and improvement opportunities. Generates strategic recommendations for the migration approach."),
    ("Questionnaire Generation", "Creates targeted, multi-choice questions specific to the migration findings. The consultant delivers these to the end client. AI pre-fills recommendations with confidence scores."),
    ("BRD Generation", "Produces a full Business Requirements Document \u2014 AS-IS state, TO-BE state, gap analysis, data mapping, and migration plan. Updates instantly based on questionnaire responses."),
    ("Implementation & Configuration", "Configures the target system (NetSuite, Dynamics 365, etc.), executes the data migration (ETL), and runs validation checks. Zero human hours on migration execution."),
    ("Testing & QA", "Automated testing, reconciliation, and proof artifacts generated for the consultant to present to the client."),
    ("Go-Live Execution", "Cutover and final verification. The consultant handles UAT and training with the client."),
]

for label, desc in capabilities:
    bullet(desc, bold_prefix=f"{label} \u2014 ")

doc.add_heading("What the Consultant Still Handles", level=2)

para(
    "The consultant\u2019s role becomes focused on the high-value, human-centric parts of the "
    "engagement. This is lightweight but essential:"
)

consultant_tasks = [
    "Running the initial discovery call and capturing client context",
    "Providing business context and institutional knowledge to Source",
    "Reviewing Source AI\u2019s outputs (BRD, migration plan, questionnaires)",
    "Delivering questionnaires to the end user and facilitating sign-off",
    "Confirming the final BRD with the end user",
    "User training, UAT, and go-live handover",
]
for task in consultant_tasks:
    bullet(task)

doc.add_heading("What the End Client Does", level=2)

para(
    "The end client\u2019s involvement is minimal \u2014 they answer questionnaires, confirm their "
    "business requirements, sign off on the TO-BE state, and accept the final migration. "
    "They interact through a branded portal under the consultancy\u2019s identity."
)

doc.add_heading("The Step-by-Step Workflow", level=2)

para(
    "In practice, a Source-powered migration follows ten steps. The consultant touches "
    "steps 1, 5, and 9. The end client touches steps 6 and 10. Source handles everything else:"
)

add_table(
    ["Step", "Owner", "Action"],
    [
        ["1", "Consultant", "Client call & notes \u2014 runs normal discovery session"],
        ["2", "Source AI", "Connect & scan \u2014 read-only access to legacy system"],
        ["3", "Source AI", "AI questionnaire \u2014 generates targeted questions from scan findings"],
        ["4", "Source AI", "Draft BRD \u2014 full requirements document with migration plan"],
        ["5", "Consultant", "Consultant review \u2014 reviews outputs, adds context"],
        ["6", "Customer", "Customer sign-off \u2014 confirms requirements and TO-BE state"],
        ["7", "Source AI", "AI migration \u2014 implementation, configuration, data migration"],
        ["8", "Source AI", "Testing & QA \u2014 automated validation and proof artifacts"],
        ["9", "Consultant", "Go-live & handover \u2014 UAT, training, cutover"],
        ["10", "Customer", "Happy customer"],
    ],
    col_widths=[0.5, 1.3, 4.9],
)

para(
    "The entire process takes under 21 days end-to-end, with roughly 12 hours of consultant "
    "time across the engagement. No offshore developers. No solutions architect. No months "
    "of BRD revision cycles."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# LIVE DEMO: WHAT SOURCE ACTUALLY PRODUCES
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("What Source Actually Produces: A Live Walkthrough", level=1)

para(
    "In live demonstrations to prospective partners, Source AI\u2019s CEO Liam Fuller walks through "
    "the system using a real QuickBooks environment with over $33 million in transaction data "
    "\u2014 an intentionally complex test instance. The demo shows what happens after a consultant "
    "uploads discovery call notes and gives Source read-only API access to the client\u2019s system."
)

para(
    "Source begins by scanning the QuickBooks instance. Within minutes, it identifies the "
    "system\u2019s characteristics: deferred revenue patterns, opening balance issues, duplicate "
    "vendors, and a full AP/AR breakdown. In one demonstration, Source surfaced a critical "
    "finding that would typically take a solutions architect days to discover: all physical "
    "product revenue from the client\u2019s Shopify connector was flowing through a discount income "
    "account instead of proper sales revenue accounts. In Source\u2019s words: \u201cRevenue reported in "
    "QuickBooks is fundamentally misstated at the gross revenue line.\u201d This is the kind of "
    "finding that changes the entire migration strategy \u2014 and Source catches it automatically."
)

para("From the scan, Source generates a comprehensive set of outputs:")

demo_outputs = [
    "An executive summary with an overall data readiness assessment, account counts, vendor and customer totals, and critical findings",
    "A company profile assembled from the discovery conversation and the system scan \u2014 covering channel mix (e-commerce, retail, wholesale), supply chain, and key stakeholders",
    "A month-end close analysis comparing what the client stated in conversation against what the data actually shows",
    "A migration strategy with phasing recommendations and data mapping decisions (for example: \u201cdo not migrate this data \u2014 build it from the Shopify integration instead\u201d)",
    "Auto-generated questions for the next discovery call, each with impact assessments and the responsible stakeholder identified",
    "Option A / Option B decision points with Source\u2019s default recommendation and confidence score",
]
for output in demo_outputs:
    bullet(output)

para(
    "All of this is produced from a single system scan and a set of discovery call notes. "
    "The consultant can walk into their next call with the client armed with a document that "
    "would have taken a solutions architect a week or more to produce manually. After the "
    "second call, the questionnaire responses flow back into Source, and the BRD is updated "
    "automatically. Once the client signs off on the BRD, Source begins the migration execution."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# CASE STUDY
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Case Study: Wild Tech \u2014 QuickBooks to NetSuite in 13 Days", level=1)

para(
    "Source AI\u2019s first completed migration was delivered in partnership with Wild Tech, a "
    "NetSuite Solution Provider based in Australia. The engagement involved a mid-market "
    "retail client with approximately 120 employees, migrating from QuickBooks to NetSuite."
)

add_table(
    ["Metric", "Value"],
    [
        ["Migration time", "13 days"],
        ["Total cost to consultancy", "\u20ac11,000 (fixed price)"],
        ["Cost savings vs. traditional", "67%"],
        ["Consultant hours", "~12 hours"],
        ["Offshore developers needed", "Zero"],
    ],
    col_widths=[2.5, 4],
)

doc.add_heading("Timeline Breakdown", level=2)

add_table(
    ["Phase", "Days", "Executed By"],
    [
        ["System scan", "1", "Source AI"],
        ["Analysis", "2", "Source AI"],
        ["BRD generation", "2", "Source AI"],
        ["Configuration", "4", "Source AI"],
        ["Migration & QA", "3", "Source AI"],
        ["Go-live", "1", "Consultant"],
    ],
    col_widths=[2.5, 1, 2.5],
)

doc.add_heading("End Client Outcomes", level=2)

para(
    "Beyond the speed and cost improvements for the consultancy, the end client saw "
    "immediate operational improvements after migrating to NetSuite:"
)

add_table(
    ["Area", "Before (QuickBooks)", "After (NetSuite)"],
    [
        ["Month-end close", "12 days", "3 days"],
        ["Inventory sync", "Manual / daily", "Real-time"],
        ["Channel visibility", "Spreadsheets", "Unified dashboard"],
        ["International support", "Not possible", "Multi-currency ready"],
        ["Reporting", "48hr manual exports", "Live dashboards"],
        ["Audit trail", "Fragmented logs", "Fully automated"],
    ],
    col_widths=[1.8, 2.2, 2.5],
)

doc.add_heading("Partner Testimonial", level=2)

quote(
    "We quoted \u20ac14K fixed-price and delivered in under three weeks. Our old model would have "
    "been \u20ac35K over four months with an offshore team. Source handled the technical execution "
    "\u2014 we kept the client relationship and made 50%+ margins.",
    "Managing Director, Wild Tech \u2014 NetSuite Solution Provider, Australia"
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# SUPPORTED SYSTEMS
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Supported Migration Paths", level=1)

para(
    "Source currently supports the most common mid-market migration paths. These cover the "
    "majority of ERP migrations in the market \u2014 NetSuite alone sees approximately 22,000 new "
    "customers per year, and 55% of those come from QuickBooks."
)

doc.add_heading("Live Today", level=2)

add_table(
    ["From", "To", "Integrations Supported"],
    [
        ["QuickBooks", "NetSuite", "Shopify, HubSpot, Stripe, + 100s more"],
        ["Xero", "NetSuite", "Stripe, PayPal, + 100s more"],
        ["QuickBooks", "Dynamics 365 / Business Central", "Salesforce, SharePoint, + 100s more"],
    ],
    col_widths=[1.5, 2.3, 3],
)

para(
    "These migrations are not standalone \u2014 mid-market companies typically have multiple "
    "systems, integrations, and workflows attached to their core accounting system. Source "
    "handles the integrations as well: Shopify e-commerce connectors, HubSpot CRM syncs, "
    "Stripe payment data, and hundreds of other common integrations."
)

para(
    "Source has found that Xero was relatively straightforward to support compared to "
    "QuickBooks, which tends to have more complex and messy data structures. The next "
    "expansion targets include MYOB (which dominates the legacy market in Australia), "
    "Acumatica, Freshbooks, Zoho Books, DEAR Systems, Reckon, Wave, and custom or legacy "
    "on-premise ERPs. The long-term vision is \u201cAnything \u2192 NetSuite.\u201d"
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# WHY PARTNER
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Why Consultancies Partner with Source AI", level=1)

para(
    "Source AI does not replace your consultancy. It makes every engagement more profitable, "
    "faster, and more scalable. The value proposition breaks down into six concrete benefits:"
)

doc.add_heading("1. Make More Money", level=2)

para(
    "Traditional mid-market migration margins sit at 20\u201330%. With Source, consultants "
    "regularly achieve 50%+ margins on migration projects. The reason is simple: AI execution "
    "eliminates the need for offshore teams, solutions architects, and months of billable "
    "overhead. The consultancy keeps the commercial upside while Source absorbs the delivery cost."
)

doc.add_heading("2. Happier Customers", level=2)

para(
    "Migrations that used to take 4\u20136 months now complete in under 21 days. Faster delivery "
    "means fewer revision cycles, fewer surprises, and a client who gets to their new system "
    "while the momentum from the buying decision is still fresh. Every deliverable comes with "
    "validated artifacts \u2014 the consultant presents proof, not promises."
)

doc.add_heading("3. Fixed-Price Deals", level=2)

para(
    "Because Source provides a fixed-price quote to the consultancy after the initial system "
    "scan, the consultant knows exactly what the migration will cost before quoting the client. "
    "This makes fixed-price proposals possible \u2014 something most consultancies avoid because "
    "of scope creep risk. With Source, scope creep is caught by the AI during the system scan, "
    "not discovered mid-build."
)

doc.add_heading("4. Scale Without Hiring", level=2)

para(
    "A consultancy can grow its migration practice without recruiting, onboarding, or managing "
    "offshore development teams. Source runs autonomously. There are zero new hires needed to "
    "double or triple migration throughput."
)

doc.add_heading("5. Win More Deals", level=2)

para(
    "When competitors are quoting $30K\u2013$40K for a migration, a Source-powered consultancy can "
    "quote $8K\u2013$15K and still keep 50%+ margins. The consultancy wins on price, wins on speed, "
    "and wins on the ability to offer a fixed-price guarantee. It is a fundamentally different "
    "competitive position."
)

doc.add_heading("6. White Label Delivery", level=2)

para(
    "Everything Source produces is delivered under the consultancy\u2019s brand. Client-facing "
    "questionnaires appear on the consultancy\u2019s domain, with the consultancy\u2019s logo, colour "
    "scheme, and consultant names as authors. No Source AI branding is visible to the end "
    "client. The consultancy comes across as a highly efficient, fast-moving firm."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# PRICING
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Pricing and Partnership Model", level=1)

para(
    "Source AI partners with ERP consultancies on a simple, transparent model. The economics "
    "are designed so that both parties make significantly more money than the traditional "
    "delivery model allows."
)

doc.add_heading("How It Works", level=2)

bold_para("The first migration is free.")

para(
    "Source provides the first migration at no cost and no commitment. This allows the "
    "consultancy to see the quality of Source\u2019s output on a real engagement before making "
    "any financial commitment."
)

para(
    "After the first migration, Source operates on a revenue share model. Source provides a "
    "fixed price to the consultancy for each migration. The consultancy marks it up however "
    "they choose and quotes the end client. The consultancy keeps the margin."
)

doc.add_heading("Example Economics", level=2)

para(
    "On a traditional $30K migration, a consultancy nets approximately $6K profit (20% margin) "
    "after accounting for sales costs, PM overhead, offshore teams, and implementation time."
)

para(
    "With Source, the consultant quotes the client $15K for the same migration. Source\u2019s "
    "fixed-price cost might be $4,500. The consultant keeps $10,500 \u2014 with roughly 12 hours "
    "of human work instead of 80\u2013130 hours. That\u2019s a 70% margin instead of 20%. The client "
    "saves $15K+ and gets it done in weeks instead of months. Everyone wins."
)

doc.add_heading("Traditional vs. Source AI Comparison", level=2)

add_table(
    ["", "Traditional", "With Source AI"],
    [
        ["Timeline", "6\u201312+ months", "Sub 21 days"],
        ["Typical cost to client", "$30K\u2013$40K", "$8K\u2013$15K"],
        ["BRD revisions", "3\u20135 cycles, weeks of delay", "AI-generated, instant updates"],
        ["Scope creep risk", "High \u2014 discovered mid-build", "Low \u2014 AI system scan upfront"],
        ["Offshore team ramp-up", "2\u20136 weeks onboarding", "No team needed"],
        ["Client satisfaction", "Frustrated by delays & overruns", "Fast, transparent, fewer surprises"],
        ["Pricing model", "T&M \u2014 unpredictable bills", "Fixed price \u2014 locked quote upfront"],
    ],
    col_widths=[1.8, 2.2, 2.7],
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# MARKET CONTEXT
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Market Context: AI and the Future of ERP", level=1)

para(
    "The ERP industry is at an inflection point. AI has been, in the words of one senior "
    "Australian ERP consultant, \u201ca real game changer\u201d in the last 12\u201318 months \u2014 both in "
    "how consulting work is delivered and in what the ERP vendors themselves are building."
)

para(
    "The most significant development is around Model Context Protocol (MCP) connectors, "
    "which allow companies to bring their own AI agent and connect it directly to an ERP "
    "database. Microsoft Dynamics is pushing MCP aggressively, with impressive live demos "
    "already in the market. NetSuite has a broad product portfolio but its AI timelines \u2014 "
    "targeting late 2026 \u2014 feel slow to many in the industry. Other vendors are attempting "
    "to build AI directly into their platforms with varying degrees of success, and there is "
    "a growing concern about technical debt in features that may be irrelevant within six months."
)

para(
    "The consensus among forward-thinking ERP consultants is that the traditional SaaS "
    "licensing model is going to be \u201cchallenged in a pretty significant way.\u201d The question "
    "being asked is: if an IT director with basic prompting skills can spin up an ERP system "
    "in two to three years, where do these platforms find their value? The vendors that do not "
    "move fast enough will be left behind."
)

para(
    "There is nuance, of course. Some functional areas \u2014 particularly warehousing, WMS, and "
    "shop-floor operations \u2014 still require traditional interfaces. But for finance teams, "
    "reporting, sales, customer service, and most back-office functions, the industry expects "
    "significant change in the near term."
)

para(
    "Source AI\u2019s bet is straightforward: regardless of how ERP platforms evolve, companies "
    "will continue to need to get off their legacy systems. That migration process is painful "
    "today and will remain a bottleneck even as the destination systems become smarter. Source "
    "is building the infrastructure layer that powers every mid-market ERP migration."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# COMPETITIVE DIFFERENTIATION
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Competitive Differentiation", level=1)

para(
    "The ERP migration automation space is emerging. One competitor that comes up in "
    "conversations is Trove, which takes a fundamentally different approach:"
)

add_table(
    ["", "Source AI", "Trove"],
    [
        ["Go-to-market", "Partners WITH consultancies", "Goes DIRECT to end clients"],
        ["Consultant role", "Keeps the relationship & commercial ownership", "Attempts to remove consultants entirely"],
        ["Philosophy", "Consultants + AI = best outcome", "AI replaces the consultant"],
    ],
    col_widths=[1.5, 2.6, 2.6],
)

para(
    "Source\u2019s view is that AI is not yet at the point where it can replace the human consultant "
    "entirely for high-stakes, costly ERP migrations. Clients making a major business decision "
    "\u2014 a decision that affects their entire operation \u2014 want a human relationship. They want "
    "someone they can call, someone who understands their business context, and someone who takes "
    "accountability. Source makes that consultant dramatically more effective, but it does not "
    "try to remove them from the equation."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# COMPLEXITY
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Handling Complexity", level=1)

para(
    "A common question from prospective partners is whether Source can handle complex clients "
    "\u2014 not just core finance, but businesses with wholesale distribution, warehouse management, "
    "manufacturing, or global multi-subsidiary operations."
)

para(
    "Source handles complexity in terms of data depth. If a company has one central system "
    "with deep data \u2014 accounts payable, accounts receivable, complex integrations, workflows, "
    "and customisations attached to it \u2014 Source handles that well. The complexity boundary "
    "today is the number of completely disparate source systems, not the depth or complexity "
    "of the data within a single system. A QuickBooks instance with Shopify, HubSpot, Stripe, "
    "and complex revenue recognition is well within scope. Four completely separate legacy ERPs "
    "that need to be consolidated would require a phased approach."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# THE TEAM
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("The Team", level=1)

para(
    "Source AI is built by a team that combines AI engineering capability with deep ERP "
    "domain expertise. The company is based in San Francisco, California."
)

add_table(
    ["Name", "Role", "Background"],
    [
        ["Liam Fuller", "CEO", "Built Source\u2019s autonomous migration engine. AI infrastructure & multi-agent systems."],
        ["Yoan Gabison", "CTO", "10+ years in NetSuite & ERP. Full-stack AI architect who designed Source\u2019s core platform."],
        ["Shane Fuller", "Founding Engineer", "Former Cybersecurity IT Director. Enterprise security, infrastructure & compliance."],
        ["Brian Kelleher", "AI Engineer", "ML pipelines & data migration. Builds the models that power Source\u2019s automation."],
        ["Jiri Pucs", "AI Engineer", "Backend systems, API architecture & integration pipelines."],
    ],
    col_widths=[1.5, 1.3, 3.9],
)

para(
    "The team also includes three senior ERP consultants \u2014 Sean Gillespie, Matthew Rodgers, "
    "and Bryan McCutchan \u2014 who bring hands-on experience in ERP strategy, implementation, "
    "and client delivery. Several have backgrounds at major ERP consultancies (SAP, Dynamics, "
    "NetSuite practices), and one is an advisor who previously exited an ERP consultancy to "
    "Microsoft. An additional 10+ internal expertise consultants support the team."
)

para(
    "The team\u2019s approach to building Source itself reflects the AI-native philosophy: with "
    "just two core engineers, Source produces millions of lines of code, spending tens of "
    "thousands of euros per month on AI compute credits. The CTO and CEO work at near-parity "
    "in terms of code output \u2014 a dynamic that would have been impossible 18 months ago."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# CURRENT PARTNERS & TRACTION
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Current Partners and Traction", level=1)

para("Source AI has early traction with several signed partnerships:")

bullet("Wild Tech \u2014 NetSuite Solution Provider, Australia. First completed migration (13 days, \u20ac11K).", bold_prefix="")
bullet("Bring IT \u2014 the 2nd largest NetSuite consultancy in the USA, with 300+ ERP experts across three continents. Partnership to automate 85%+ of migration contracts under $50K.", bold_prefix="")
bullet("AppWarp \u2014 Australian NetSuite consultancy.", bold_prefix="")
bullet("3 signed migration contracts at $15K each with consultancies.", bold_prefix="")

para(
    "Source is actively seeking early consultancy partners. The first migration is free of "
    "charge, with no commitment required. The goal is to onboard the top 50\u2013100 mid-market "
    "ERP consultancies and reach 10 QuickBooks \u2192 NetSuite migrations per month."
)

add_hr()

# ═════════════════════════════════════════════════════════════════════════════
# THE VISION
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("The Vision", level=1)

para(
    "The consulting industry was built on the back of ERP vendors \u2014 SAP, NetSuite, Microsoft "
    "Dynamics \u2014 who could not scale into the hundreds of thousands of mid-market businesses "
    "across manufacturing, healthcare, education, local government, retail, and more. They "
    "outsourced implementation to consultancies in exchange for aggressive license kickbacks."
)

para(
    "Source AI collapses that delivery chain. The consultant becomes the relationship owner "
    "and quality controller. The AI does the work. The client gets a faster, cheaper, better "
    "migration. Everyone wins."
)

para(
    "Source is not a tool. It fundamentally changes how consulting works in the age of AI."
)

for _ in range(2):
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# KEY NUMBERS
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Key Numbers at a Glance", level=1)

add_table(
    ["", "Value"],
    [
        ["Migration automation", "80% of the work handled by AI"],
        ["Migration timeline (with Source)", "Sub 21 days"],
        ["Traditional timeline", "4\u20136+ months"],
        ["Cost to client (with Source)", "$8K\u2013$15K"],
        ["Traditional cost to client", "$30K\u2013$40K"],
        ["Consultant margins (with Source)", "50%+"],
        ["Traditional margins", "~20%"],
        ["Case study: Wild Tech", "13 days, \u20ac11K, 67% savings"],
        ["Source AI work share", "85% autonomous"],
        ["First migration pricing", "FREE"],
        ["Team", "5 core + 3 consultants + 10 extended"],
        ["HQ", "San Francisco, CA"],
    ],
    col_widths=[2.8, 3.9],
)

# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════

output_path = "/Users/liam/projects/source-ai-pitch/Source-AI-Pitch-Deck-Content.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
